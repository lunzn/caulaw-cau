#!/usr/bin/env python3
"""
Export teacher research summary to Word document.
Usage: python3 export-summary.py <teacherId> [--type=all|papers|patents|collab] [--output=/tmp/summary.docx]
"""
import json, sys, os, time
from urllib.request import urlopen
from urllib.error import URLError
from urllib.parse import quote

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("缺少 python-docx，正在安装...")
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "-q"], check=True)
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

teacher_id     = sys.argv[1] if len(sys.argv) > 1 else ""
export_type    = "all"
region_filter  = ""
output_path    = f"/tmp/research_summary_{teacher_id}_{int(time.time())}.docx"
for arg in sys.argv[2:]:
    if arg.startswith("--type="):   export_type   = arg[7:]
    if arg.startswith("--output="): output_path   = arg[9:]
    if arg.startswith("--region="): region_filter = arg[9:]

if not teacher_id:
    print("用法: python3 export-summary.py <teacherId> [--type=all|papers|patents|collab]")
    sys.exit(1)

base = os.environ.get("SCHOOL_SERVER_URL", "http://school-server:3002")

def fetch(url):
    with urlopen(url) as r:
        return json.loads(r.read())

try:
    teacher      = fetch(f"{base}/api/teachers/{teacher_id}")["data"]
    paper_qs     = f"?region={quote(region_filter)}" if region_filter else ""
    patent_qs    = f"?region={quote(region_filter)}" if region_filter else ""
    papers       = fetch(f"{base}/api/teachers/{teacher_id}/papers{paper_qs}")["data"]
    patents      = fetch(f"{base}/api/teachers/{teacher_id}/patents{patent_qs}")["data"]
    all_teachers = fetch(f"{base}/api/teachers")["data"]
except URLError as e:
    print(f"API调用失败: {e}")
    sys.exit(1)

name = teacher.get("name", teacher_id)
dept = teacher.get("department", "")
title_s = teacher.get("title", "")

doc = Document()

# ── 页面设置 ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(3)
section.top_margin  = section.bottom_margin = Cm(2.5)

# ── 标题样式辅助 ──────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_bold_para(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p

# ── 文档标题 ──────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
region_label = f"（{region_filter}地区）" if region_filter else ""
run = title_para.add_run(f"{name} {title_s} 科研成果报告{region_label}")
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph(f"生成日期：{time.strftime('%Y年%m月%d日')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── 教师简介 ──────────────────────────────────────────────────────────────────
if export_type in ("all", "profile"):
    add_heading(doc, "一、基本信息", level=1)
    info = [
        ("姓名", name),
        ("职称", title_s),
        ("院系", dept),
        ("研究方向", teacher.get("research_areas") or "—"),
        ("办公室", teacher.get("office") or "—"),
        ("邮箱", teacher.get("email") or "—"),
    ]
    for label, val in info:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}：").bold = True
        p.add_run(val)
    doc.add_paragraph()

# ── 论文统计 ──────────────────────────────────────────────────────────────────
if export_type in ("all", "papers"):
    add_heading(doc, "二、学术论文" if export_type == "all" else "学术论文", level=1)

    total_cite = sum(p["citation_count"] for p in papers)
    from collections import Counter
    years_cnt   = Counter(p["year"] for p in papers)
    regions_cnt = Counter(p["region"].split("（")[0] for p in papers)

    add_bold_para(doc, f"论文总数：{len(papers)} 篇  总被引：{total_cite} 次  均被引：{total_cite/max(len(papers),1):.1f} 次/篇")
    doc.add_paragraph()

    add_bold_para(doc, "▸ 年份分布", size=10)
    for yr in sorted(years_cnt.keys(), reverse=True):
        doc.add_paragraph(f"  {yr}年：{years_cnt[yr]} 篇", style="List Bullet")

    add_bold_para(doc, "▸ 发表地区", size=10)
    for reg, cnt in regions_cnt.most_common():
        doc.add_paragraph(f"  {reg}：{cnt} 篇", style="List Bullet")

    doc.add_paragraph()
    add_bold_para(doc, "▸ 代表性论文（被引Top20）", size=10)
    top_papers = sorted(papers, key=lambda x: x["citation_count"], reverse=True)[:20]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ["序号", "题目", "期刊/会议", "年份/被引"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
    for i, p in enumerate(top_papers, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = p["title"]
        row[2].text = p["journal"]
        row[3].text = f"{p['year']}年，{p['citation_count']}次"
    doc.add_paragraph()

# ── 知识产权 ─────────────────────────────────────────────────────────────────
if export_type in ("all", "patents"):
    add_heading(doc, "三、知识产权" if export_type == "all" else "知识产权", level=1)

    from collections import Counter
    types_cnt = Counter(p["type"] for p in patents)
    add_bold_para(doc, f"合计：{len(patents)} 项  " + "  ".join(f"{t}:{c}项" for t, c in types_cnt.most_common()))
    doc.add_paragraph()

    for ip_type in ["发明专利", "实用新型", "软件著作权"]:
        items = [p for p in patents if p["type"] == ip_type]
        if not items: continue
        add_bold_para(doc, f"▸ {ip_type}（{len(items)} 项）", size=10)
        for i, p in enumerate(sorted(items, key=lambda x: x["year"], reverse=True), 1):
            cert = f"  证书号：{p['cert_number']}" if p.get("cert_number") else ""
            reg  = f"  [{p['region']}]" if "港澳" in p.get("region", "") else ""
            doc.add_paragraph(f"  {i}. {p['title']}（{p['year']}年）{cert}{reg}", style="List Bullet")
        doc.add_paragraph()

# ── 潜在合作者 ────────────────────────────────────────────────────────────────
if export_type in ("all", "collab"):
    add_heading(doc, "四、校内潜在合作者" if export_type == "all" else "校内潜在合作者", level=1)

    own_areas = set((teacher.get("research_areas") or "").split(","))
    collabs = []
    for t in all_teachers:
        if t["id"] == teacher_id: continue
        t_areas = (t.get("research_areas") or "") + " " + t.get("department", "")
        for area in own_areas:
            area = area.strip()
            if area and area in t_areas:
                collabs.append((t, area))
                break

    if collabs:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for cell, txt in zip(hdr, ["姓名", "职称", "院系", "邮箱"]):
            cell.text = txt
            cell.paragraphs[0].runs[0].bold = True
        for t, area in collabs:
            row = table.add_row().cells
            row[0].text = t["name"]
            row[1].text = t.get("title", "—")
            row[2].text = t.get("department", "—")
            row[3].text = t.get("email", "—")
    else:
        doc.add_paragraph("暂无精确匹配的校内合作者（建议通过信电学院官网扩展搜索）")
    doc.add_paragraph()

# ── 保存 ──────────────────────────────────────────────────────────────────────
doc.save(output_path)
print(f"✅ Word 文档已生成")
print(f"FILE:{output_path}")
